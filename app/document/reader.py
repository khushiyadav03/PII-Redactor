from dataclasses import dataclass
from typing import List, Tuple

import docx
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class RunSpan:
    # Paragraph ke logical text mein ek run ka position.
    run_index: int

    # Logical text mein run kahan se start hota hai.
    start: int

    # Logical text mein run kahan end hota hai.
    # End exclusive hai.
    end: int


@dataclass
class LogicalParagraph:
    # Original python-docx Paragraph object.
    # Writer isi object ke runs ko modify karega.
    paragraph: Paragraph

    # Saare runs ko join karke bana hua complete paragraph text.
    text: str

    # Logical text aur original runs ke beech mapping.
    run_spans: List[RunSpan]

    # Paragraph document ke kis part se aaya hai.
    location: str


def _build_logical_text(
    paragraph: Paragraph,
) -> Tuple[str, List[RunSpan]]:
    # Ek paragraph ke saare runs ko join karke
    # ek continuous logical string banate hain.

    text_parts = []
    spans: List[RunSpan] = []

    # Ye logical string mein current character position track karta hai.
    cursor = 0

    for idx, run in enumerate(paragraph.runs):
        # run.text None ho sakta hai, isliye empty string fallback.
        run_text = run.text or ""

        # Empty run ka text mein koi contribution nahi hai.
        if run_text == "":
            continue

        # Current run logical text mein yahan se start hota hai.
        start = cursor

        # Current run yahan tak occupy karega.
        end = cursor + len(run_text)

        # Run ki position save kar lo.
        spans.append(
            RunSpan(
                run_index=idx,
                start=start,
                end=end,
            )
        )

        # Run ka actual text collect karo.
        text_parts.append(run_text)

        # Cursor ko next position par move karo.
        cursor = end

    # Saare run texts ko ek continuous string mein join karo.
    return "".join(text_parts), spans


def iter_logical_paragraphs(
    paragraphs: List[Paragraph],
    location_prefix: str,
) -> List[LogicalParagraph]:

    # Normal Paragraph objects ko LogicalParagraph objects mein convert karta hai.
    result = []

    for i, paragraph in enumerate(paragraphs):
        # Paragraph ka combined text aur run mapping banao.
        text, spans = _build_logical_text(paragraph)

        # Completely empty paragraphs ko ignore karo.
        if text.strip() == "":
            continue

        # Logical paragraph ko result mein add karo.
        result.append(
            LogicalParagraph(
                paragraph=paragraph,
                text=text,
                run_spans=spans,
                location=f"{location_prefix}:para{i}",
            )
        )

    return result


def _iter_table_paragraphs(
    table: Table,
    location_prefix: str,
) -> List[LogicalParagraph]:

    # Table ke andar har cell ke paragraphs process karta hai.
    result = []

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):

            # Current cell ka unique location.
            loc = f"{location_prefix}:row{r_idx}:cell{c_idx}"

            # Cell ke normal paragraphs process karo.
            result.extend(
                iter_logical_paragraphs(
                    cell.paragraphs,
                    loc,
                )
            )

            # Cell ke andar nested tables bhi ho sakti hain.
            # Isliye unhe recursively process karte hain.
            for nt_idx, nested_table in enumerate(cell.tables):
                result.extend(
                    _iter_table_paragraphs(
                        nested_table,
                        f"{loc}:nested_table{nt_idx}",
                    )
                )

    return result


@dataclass
class DocumentContent:
    # Original DOCX document object.
    document: DocxDocument

    # Document ke different areas ke logical paragraphs.
    body_paragraphs: List[LogicalParagraph]
    table_paragraphs: List[LogicalParagraph]
    header_paragraphs: List[LogicalParagraph]
    footer_paragraphs: List[LogicalParagraph]

    def all_paragraphs(self) -> List[LogicalParagraph]:
        # Sabhi areas ke paragraphs ko ek single list mein combine karo.
        return (
            self.body_paragraphs
            + self.table_paragraphs
            + self.header_paragraphs
            + self.footer_paragraphs
        )


def load_document(path: str) -> DocumentContent:
    # DOCX ko open karo aur uske saare important text areas
    # ko LogicalParagraph format mein convert karo.

    try:
        # python-docx DOCX file ko parse karke Document object deta hai.
        doc = docx.Document(path)

    except Exception as exc:
        # Invalid ya corrupted DOCX ke case mein
        # ek clear error raise karo.
        raise ValueError(
            f"Could not open DOCX file '{path}': {exc}"
        ) from exc

    # Main document body ke paragraphs process karo.
    body_paragraphs = iter_logical_paragraphs(
        doc.paragraphs,
        "body",
    )

    # Tables ke paragraphs alag collect karo.
    table_paragraphs: List[LogicalParagraph] = []

    for t_idx, table in enumerate(doc.tables):
        table_paragraphs.extend(
            _iter_table_paragraphs(
                table,
                f"table{t_idx}",
            )
        )

    # Headers aur footers ke paragraphs.
    header_paragraphs: List[LogicalParagraph] = []
    footer_paragraphs: List[LogicalParagraph] = []

    # Same header/footer multiple sections mein shared ho sakta hai.
    # Isliye duplicate objects ko process hone se bachate hain.
    seen_header_ids = set()
    seen_footer_ids = set()

    # DOCX ke har section ko process karo.
    for s_idx, section in enumerate(doc.sections):

        # Header aur footer dono ko same logic se process kar sakte hain.
        for header_attr, seen_set, bucket, label in [
            (
                section.header,
                seen_header_ids,
                header_paragraphs,
                "header",
            ),
            (
                section.footer,
                seen_footer_ids,
                footer_paragraphs,
                "footer",
            ),
        ]:

            # Header/footer ke underlying XML element ki identity.
            part_id = id(header_attr._element)

            # Agar same object pehle process ho chuka hai,
            # toh dobara process mat karo.
            if part_id in seen_set:
                continue

            seen_set.add(part_id)

            # Header/footer ke normal paragraphs process karo.
            bucket.extend(
                iter_logical_paragraphs(
                    header_attr.paragraphs,
                    f"{label}{s_idx}",
                )
            )

            # Header/footer ke andar tables bhi ho sakti hain.
            for t_idx, table in enumerate(header_attr.tables):
                bucket.extend(
                    _iter_table_paragraphs(
                        table,
                        f"{label}{s_idx}:table{t_idx}",
                    )
                )

    # Saari processed information ek object mein return karo.
    return DocumentContent(
        document=doc,
        body_paragraphs=body_paragraphs,
        table_paragraphs=table_paragraphs,
        header_paragraphs=header_paragraphs,
        footer_paragraphs=footer_paragraphs,
    )