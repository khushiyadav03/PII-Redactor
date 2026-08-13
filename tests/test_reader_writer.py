"""
Hinglish: Ye sabse critical test hai - agar split-run handling galat hui
to poora PII detection unreliable ho jayega. "Rahul Sharma" ko "Rah"+"ul "+
"Sharma" runs mein todkar test karte hain (jaisa Word real documents mein
karta hai), aur verify karte hain ki:
  1. Logical text sahi reconstruct hota hai.
  2. Redaction sahi run mein apply hoti hai aur original text kahin bachta nahi.
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.document.reader import load_document
from app.document.writer import apply_redactions

SAMPLE = os.path.join(os.path.dirname(__file__), "split_run_sample.docx")


def test_logical_text_reconstructs_across_split_runs():
    content = load_document(SAMPLE)
    assert len(content.body_paragraphs) == 1
    lp = content.body_paragraphs[0]
    assert lp.text == "Rahul Sharma works at Infosys."


def test_redaction_across_split_runs_removes_original_text():
    content = load_document(SAMPLE)
    lp = content.body_paragraphs[0]
    start = lp.text.index("Rahul Sharma")
    end = start + len("Rahul Sharma")
    apply_redactions(lp, [(start, end, "Alex Morgan")])

    # Hinglish: naya text rebuild karke check karo original "Rahul" ya
    # "Sharma" kahin bhi (kisi bhi run mein) bacha na ho.
    full_text_after = "".join(r.text for r in lp.paragraph.runs)
    assert "Rahul" not in full_text_after
    assert "Sharma" not in full_text_after
    assert "Alex Morgan" in full_text_after
    assert "Infosys" in full_text_after  # unrelated text untouched


def test_redaction_preserves_surrounding_text_in_same_run():
    content = load_document(SAMPLE)
    lp = content.body_paragraphs[0]
    start = lp.text.index("Infosys")
    end = start + len("Infosys")
    apply_redactions(lp, [(start, end, "Acme Corp")])
    full_text_after = "".join(r.text for r in lp.paragraph.runs)
    assert "Acme Corp" in full_text_after
    assert full_text_after.endswith(".")  # trailing period preserved


def test_load_document_on_real_prospectus_does_not_crash():
    real_path = os.path.join(
        os.path.dirname(__file__), "..", "samples", "Red_Herring_Prospectus.docx"
    )
    content = load_document(real_path)
    assert len(content.body_paragraphs) > 0
    assert len(content.table_paragraphs) > 0
    assert len(content.document.tables) == 76


def test_metadata_cleanup_on_save_and_reload(tmp_path):
    import docx
    from app.document.metadata import clean_core_metadata
    
    doc_path = tmp_path / "test_meta.docx"
    doc = docx.Document()
    doc.add_paragraph("Hello world.")
    
    props = doc.core_properties
    props.author = "John Doe"
    props.last_modified_by = "Jane Doe"
    props.title = "Confidential Report"
    props.comments = "This is a comment."
    props.category = "Internal"
    
    doc.save(doc_path)
    
    doc2 = docx.Document(doc_path)
    assert doc2.core_properties.author == "John Doe"
    
    clean_core_metadata(doc2)
    doc2.save(doc_path)
    
    doc3 = docx.Document(doc_path)
    assert doc3.core_properties.author == ""
    assert doc3.core_properties.last_modified_by == ""
    assert doc3.core_properties.title == ""
    assert doc3.core_properties.comments == ""
    assert doc3.core_properties.category == ""
