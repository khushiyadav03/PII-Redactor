import io
import os

import docx
import pytest
from fastapi.testclient import TestClient

from app.stream_protocol import parse_stream_response
from app.api import app

client = TestClient(app)

SAMPLE_DOCX = os.path.join(os.path.dirname(__file__), "split_run_sample.docx")


def _post_redact(file_path: str, filename: str = None):
    filename = filename or os.path.basename(file_path)
    with open(file_path, "rb") as f:
        return client.post(
            "/redact",
            files={
                "file": (
                    filename,
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )


def _make_pii_docx(path: str) -> None:
    doc = docx.Document()
    doc.add_paragraph("Name: Rahul Sharma")
    doc.add_paragraph("Email: rahul.sharma@example.com")
    doc.add_paragraph("Phone: +91 9876543210")
    doc.add_paragraph("Date of Birth: 15/08/1990")
    doc.add_paragraph("Registered Office: 42 MG Road, Bengaluru, Karnataka 560001")
    doc.add_paragraph("PAN: ABCDE1234F")
    doc.add_paragraph("Aadhaar number is 2345 6789 0123 for KYC verification.")
    doc.save(path)


def _make_clean_docx(path: str) -> None:
    doc = docx.Document()
    doc.add_paragraph("This document contains no personal identifiers.")
    doc.add_paragraph("SEBI ICDR Regulations apply to this filing.")
    doc.save(path)


def test_health_endpoint():
    """Hinglish: Verify that /health endpoint is working and returns standard status."""
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_root_serves_frontend():
    """Hinglish: Verify that / route serves the index.html page."""
    response = client.get("/")
    assert response.status_code == 200
    assert "<title>PII Redactor" in response.text
    assert 'id="loaderText"' in response.text
    assert "PII Redactions" in response.text
    assert "Images Protected" in response.text


def test_redact_endpoint_with_valid_docx():
    """Hinglish: Stream response valid DOCX return karta hai."""
    if not os.path.exists(SAMPLE_DOCX):
        pytest.skip("split_run_sample.docx not found in tests folder.")

    response = _post_redact(SAMPLE_DOCX)
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/x-pii-redactor-stream")

    metrics, docx_bytes, error = parse_stream_response(response.content)
    assert error is None
    assert docx_bytes.startswith(b"PK")
    assert metrics is not None
    assert metrics["paragraphs_scanned"] >= 0


def test_redact_returns_nonzero_metrics_for_pii_fixture(tmp_path):
    """Hinglish: Known PII fixture par backend ke actual counters non-zero hone chahiye."""
    pii_doc = tmp_path / "pii_sample.docx"
    _make_pii_docx(str(pii_doc))

    response = _post_redact(str(pii_doc), filename="pii_sample.docx")
    assert response.status_code == 200

    metrics, docx_bytes, error = parse_stream_response(response.content)
    assert error is None
    assert docx_bytes.startswith(b"PK")
    assert metrics["text_redactions_applied"] > 0


def test_redact_returns_zero_redactions_for_clean_document(tmp_path):
    """Hinglish: Bina PII document par redaction count 0 hona chahiye."""
    clean_doc = tmp_path / "clean.docx"
    _make_clean_docx(str(clean_doc))

    response = _post_redact(str(clean_doc), filename="clean.docx")
    assert response.status_code == 200

    metrics, docx_bytes, error = parse_stream_response(response.content)
    assert error is None
    assert docx_bytes.startswith(b"PK")
    assert metrics["text_redactions_applied"] == 0
    assert metrics["images_modified"] == 0


def test_redact_stream_includes_progress_events(tmp_path):
    """Hinglish: Long-processing feedback ke liye stream mein real progress lines aati hain."""
    pii_doc = tmp_path / "progress_sample.docx"
    _make_pii_docx(str(pii_doc))

    response = _post_redact(str(pii_doc), filename="progress_sample.docx")
    assert response.status_code == 200
    assert b"PROGRESS:" in response.content
    assert b"analyzing_text" in response.content


def test_redact_endpoint_rejects_non_docx():
    """Hinglish: Verify /redact endpoint rejects non-docx extension."""
    response = client.post(
        "/redact",
        files={"file": ("test.txt", io.BytesIO(b"dummy text"), "text/plain")},
    )
    assert response.status_code == 400
    assert "Only DOCX files are allowed" in response.json()["detail"]


def test_redact_endpoint_rejects_oversized_file():
    """Hinglish: Verify /redact endpoint rejects files exceeding 15MB limit."""
    huge_data = io.BytesIO(b"0" * (16 * 1024 * 1024))
    response = client.post(
        "/redact",
        files={
            "file": (
                "large.docx",
                huge_data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 413
    assert "File too large" in response.json()["detail"]


def test_redact_endpoint_rejects_malformed_docx():
    """Hinglish: Corrupt DOCX par stream ERROR line return hoti hai."""
    corrupted_data = io.BytesIO(b"garbage PK signature but not zip")
    response = client.post(
        "/redact",
        files={
            "file": (
                "corrupt.docx",
                corrupted_data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    metrics, docx_bytes, error = parse_stream_response(response.content)
    assert metrics is None
    assert docx_bytes == b""
    assert error is not None
    assert "Could not open DOCX file" in error
